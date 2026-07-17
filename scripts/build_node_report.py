from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "reports" / "AI求职助手_节点报告_本机线上部署与投递模块_2026-07-15.docx"

BLUE = "2563EB"
DARK = "172033"
MUTED = "667085"
LIGHT = "F2F4F7"
GREEN = "027A48"
AMBER = "A16207"
RED = "BE123C"


def set_font(run, size=10.5, bold=False, color=DARK):
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 12.5, True, BLUE if level == 1 else DARK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        set_font(p.add_run(item))


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Inches(1.35), Inches(1.15), Inches(4.0)]
    headers = ["检查项", "结果", "证据"]
    for idx, cell in enumerate(table.rows[0].cells):
        cell.width = widths[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        shade(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(headers[idx]), 9.5, True)
    for item, status, evidence, color in rows:
        cells = table.add_row().cells
        values = [item, status, evidence]
        for idx, cell in enumerate(cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(values[idx]), 9.2, idx == 1, color if idx == 1 else DARK)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)

normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.15

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("AI JOB ASSISTANT  |  PROJECT NODE REPORT"), 8.5, True, MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("本报告记录节点完成情况、失败证据与后续工作"), 8, False, MUTED)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(4)
set_font(title.add_run("AI 求职助手项目节点报告"), 24, True, DARK)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
set_font(subtitle.add_run("节点：BOSS/实习僧投递模块 + 3K 薪资门槛 + 本机 Cloudflare Tunnel 线上部署"), 11.5, True, BLUE)

meta = doc.add_table(rows=3, cols=2)
meta.autofit = False
for row, (label, value) in zip(meta.rows, [
    ("报告日期", "2026-07-15"),
    ("项目位置", r"C:\Users\33387\Documents\1\ai-job-assistant"),
    ("当前公网地址", "https://fashion-match-wanted-pin.trycloudflare.com"),
]):
    row.cells[0].width = Inches(1.35)
    row.cells[1].width = Inches(5.15)
    for cell in row.cells:
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade(row.cells[0], LIGHT)
    set_font(row.cells[0].paragraphs[0].add_run(label), 9.5, True)
    set_font(row.cells[1].paragraphs[0].add_run(value), 9.5)

add_heading(doc, "1. 本节点已完成", 1)
add_bullets(doc, [
    "将 BOSS 直聘和实习僧投递能力集成到原 AI Job Assistant，没有新建独立软件。",
    "新增“投递执行”模块：平台登录检查、岗位/简历选择、任务预览、一次性确认令牌、执行状态和投递记录同步。",
    "BOSS 直聘真实页面已验证：账号处于登录状态；岗位详情投递入口为“立即沟通”，页面选择器包含 .btn-startchat。",
    "实习僧真实页面已验证：账号处于登录状态；岗位详情投递入口为“投个简历”，页面选择器包含 .resume_apply.com_res。",
    "“敏思跃动”已按公司全名精确加入黑名单；后端创建投递任务时会硬拦截黑名单公司。",
    "最低月薪门槛固定读取用户中心当前值 3000 元；无薪、薪资面议、缺少最低薪资或低于 3000 元的岗位无法进入自动投递任务。",
    "新增本机 Cloudflare Quick Tunnel 部署，前端和后端通过同一 HTTPS 公网地址访问。",
    "新增 Basic Auth 访问保护；未认证请求返回 401，正确账号可访问页面和 API。",
    "新增线上启动、停止和状态脚本，并保存公网地址与进程状态。",
])

add_heading(doc, "2. 验证结果", 1)
add_status_table(doc, [
    ("后端自动化测试", "通过", "6 passed；覆盖幂等提交、黑名单和低薪岗位拦截", GREEN),
    ("前端类型检查", "通过", "TypeScript --noEmit 无错误", GREEN),
    ("生产构建", "通过", "Next.js 15 构建成功，包含 Basic Auth Middleware", GREEN),
    ("公网未认证访问", "通过", "HTTP 401", GREEN),
    ("公网认证访问", "通过", "HTTP 200，页面加载 12 个模块", GREEN),
    ("公网后端 API", "通过", "/api/health 返回 status=ok", GREEN),
    ("Cloudflare 客户端", "通过", "v2026.7.1；Windows 数字签名 Valid；签名者 Cloudflare, Inc.", GREEN),
    ("真实外部投递", "未执行", "没有在本节点向任何真实岗位点击最终提交", AMBER),
])

add_heading(doc, "3. 软件当前具备的内容", 1)
add_bullets(doc, [
    "用户中心：个人信息、目标岗位、城市、最低薪资、学历经验和公司规模偏好。",
    "简历管理：文本编辑、版本管理、主简历、PDF/DOCX/TXT/Markdown 上传解析。",
    "职位管理：岗位录入、城市/薪资/学历/经验/公司规模筛选、来源链接保存。",
    "黑名单：精确或包含匹配，职位搜索和投递任务双重拦截。",
    "JD 匹配：透明评分、关键词覆盖、缺口和建议。",
    "简历优化：生成草稿，用户确认后应用，不自动虚构经历。",
    "沟通内容：按岗位和简历生成打招呼话术，支持审核。",
    "投递执行：BOSS 直聘、实习僧登录检测、任务预览、确认后页面操作、结果核验。",
    "投递记录与数据统计：状态时间线、跟进时间、漏斗和渠道统计。",
    "AI 模拟面试：连续追问、压力模式、录音/文件转写、本地 Whisper、评分报告和历史记录。",
    "线上访问：Cloudflare HTTPS 临时地址 + 单用户密码保护。",
])

add_heading(doc, "4. 是否可以做到自动化投递简历", 1)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
set_font(p.add_run("可以做到的程度："), 10.5, True, GREEN)
set_font(p.add_run("系统已经具备“人工登录/验证码 -> 自动读取页面 -> 生成单条预览 -> 用户确认 -> 自动点击投递/沟通 -> 检查成功标识 -> 写入投递记录”的技术链路。"))
add_bullets(doc, [
    "支持 BOSS 直聘单条“立即沟通”和实习僧单条“投个简历”。",
    "支持黑名单、最低 3000 元薪资和岗位来源域名硬校验。",
    "支持一次性确认令牌和重复提交拦截，避免同一任务重复投递。",
    "不支持绕过登录、验证码、安全验证或平台风控。",
    "不建议无人值守高频批量投递；目前设计为逐条预览、逐条确认。",
])

add_heading(doc, "5. 未完成事项", 1)
add_bullets(doc, [
    "尚未使用一个真实岗位执行最终“确认并投递”，因此真实投递成功闭环还没有生产证据。",
    "实习僧点击“投个简历”后可能出现简历选择弹窗；弹窗字段和最终成功提示需在一次用户批准的真实投递中验证。",
    "BOSS 直聘“立即沟通”后可能出现简历选择或打招呼流程；需要一次真实投递验证是否需要二次确认。",
    "Cloudflare 当前为 Quick Tunnel，地址会在 Tunnel 重启后变化，没有固定域名和可用性保证。",
    "本机关闭、休眠、代理 127.0.0.1:7990 停止、Edge 退出或远程调试未授权时，线上页面或投递执行会不可用。",
    "当前只使用 Basic Auth 单用户认证，不具备多用户、找回密码、审计账号或角色权限。",
])

add_heading(doc, "6. 本节点失败记录与处理", 1)
add_status_table(doc, [
    ("GitHub curl 下载", "失败后恢复", "官方连接 22 秒后重置；改为分段下载并验证数字签名", AMBER),
    ("winget 安装", "失败后终止", "进程长时间空等且无网络连接；已安全终止", AMBER),
    ("首次启动脚本", "失败后修复", "PowerShell 执行策略阻止；改用单次进程 Bypass，不修改系统策略", AMBER),
    ("首次 Tunnel 请求", "失败后修复", "cloudflared 未继承本机 7990 代理导致 API 超时；已配置出站代理", AMBER),
    ("日志读取", "失败后修复", "Windows 文件锁导致 ReadAllText 抛错；已改共享读取", AMBER),
    ("状态脚本编码", "失败后修复", "PowerShell 5 误解析无 BOM 中文；已改为 ASCII 状态文本", AMBER),
])

add_heading(doc, "7. 后续需要做什么", 1)
add_bullets(doc, [
    "由用户指定一个符合条件的真实岗位（非黑名单、最低月薪 >= 3000 元），执行一次受控投递验证。",
    "根据真实弹窗补齐 BOSS 直聘/实习僧的简历选择、沟通内容填写和最终成功标识。",
    "如需固定网址，登录 Cloudflare 账号创建 Named Tunnel 并绑定自有域名。",
    "如需电脑重启后自动恢复，增加 Windows 任务计划：启动后端、前端、Tunnel 与 Edge 浏览器桥接。",
    "如需多人使用，增加正式用户登录、密码哈希、会话过期、权限、审计日志和敏感数据加密。",
    "增加岗位去重、每日投递上限、公司/职位风险规则和失败重试队列。",
])

add_heading(doc, "8. 线上使用信息", 1)
add_bullets(doc, [
    "公网地址：https://fashion-match-wanted-pin.trycloudflare.com",
    "访问用户名：admin",
    "访问密码：保存在项目根目录 .env.online；请勿公开分享。",
    "启动：powershell.exe -NoProfile -ExecutionPolicy Bypass -File .\\scripts\\start-online.ps1",
    "状态：powershell.exe -NoProfile -ExecutionPolicy Bypass -File .\\scripts\\status-online.ps1",
    "停止：powershell.exe -NoProfile -ExecutionPolicy Bypass -File .\\scripts\\stop-online.ps1",
])

doc.core_properties.title = "AI 求职助手项目节点报告"
doc.core_properties.subject = "本机线上部署与招聘平台投递模块"
doc.core_properties.author = "AI Job Assistant Project"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
