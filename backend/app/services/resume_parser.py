from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from pypdf import PdfReader


class ResumeParseError(ValueError):
    pass


MAX_PDF_PAGES = 100
MAX_DOCX_ENTRIES = 2_000
MAX_DOCX_UNCOMPRESSED_BYTES = 32 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 500


def _validate_docx_archive(content: bytes) -> None:
    try:
        with ZipFile(BytesIO(content)) as archive:
            entries = archive.infolist()
    except BadZipFile as exc:
        raise ResumeParseError("DOCX 文件结构无效") from exc
    if len(entries) > MAX_DOCX_ENTRIES:
        raise ResumeParseError("DOCX 内部文件数量过多")
    total_uncompressed = sum(entry.file_size for entry in entries)
    total_compressed = sum(entry.compress_size for entry in entries)
    if total_uncompressed > MAX_DOCX_UNCOMPRESSED_BYTES:
        raise ResumeParseError("DOCX 解压后内容过大")
    if total_uncompressed / max(total_compressed, 1) > MAX_DOCX_COMPRESSION_RATIO:
        raise ResumeParseError("DOCX 压缩比异常")


def parse_resume(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    try:
        if suffix == ".pdf":
            pages = PdfReader(BytesIO(content)).pages
            if len(pages) > MAX_PDF_PAGES:
                raise ResumeParseError(f"PDF 页数不能超过 {MAX_PDF_PAGES} 页")
            text = "\n".join(page.extract_text() or "" for page in pages)
        elif suffix == ".docx":
            _validate_docx_archive(content)
            document = Document(BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_rows = [
                " | ".join(cell.text for cell in row.cells)
                for table in document.tables
                for row in table.rows
            ]
            text = "\n".join(paragraphs + table_rows)
        elif suffix in {".txt", ".md"}:
            text = _decode_text(content)
        else:
            raise ResumeParseError("仅支持 PDF、DOCX、TXT 和 Markdown 文件")
    except ResumeParseError:
        raise
    except Exception as exc:
        raise ResumeParseError(f"无法解析该文件：{exc}") from exc

    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if len(normalized) < 20:
        raise ResumeParseError("未提取到足够的简历文字，请上传可复制文字的文件")
    return normalized[:50_000]


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ResumeParseError("文本编码不受支持，请转换为 UTF-8")
